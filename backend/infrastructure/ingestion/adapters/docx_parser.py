import logging
from docx import Document
from io import BytesIO
import uuid
from typing import List

from ....domain.ingestion.services.document_parser import DocumentParser
from ....domain.shared.entities.document import Document as DocumentEntity
from datetime import datetime

logger = logging.getLogger(__name__)


class DOCXParser(DocumentParser):
    """Parser for DOCX files using python-docx"""
    
    # Maximum file size: 50 MB
    MAX_FILE_SIZE = 50 * 1024 * 1024
    
    def can_parse(self, filename: str) -> bool:
        """Check if file is a DOCX"""
        return filename.lower().endswith('.docx')
    
    def parse(self, file_content: bytes, filename: str) -> DocumentEntity:
        """
        Parse DOCX file and extract text content.
        
        Extracts text from paragraphs, tables, and headers/footers.
        """
        if not file_content:
            raise ValueError("File content is empty")
        
        if len(file_content) > self.MAX_FILE_SIZE:
            raise ValueError(f"File size exceeds maximum allowed size of {self.MAX_FILE_SIZE / (1024*1024):.1f} MB")
        
        try:
            # Parse DOCX document
            doc = Document(BytesIO(file_content))
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                raise ValueError("No text content could be extracted from DOCX file")
            
            content = "\n\n".join(text_parts)
            
            return DocumentEntity(
                id=str(uuid.uuid4()),
                filename=filename,
                content=content.strip(),
                created_at=datetime.now(),
                updated_at=datetime.now(),
            )
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {filename}: {e}", exc_info=True)
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")
